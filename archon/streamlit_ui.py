"""
Construction Submittal Document Generator
Interactive tool for creating construction submittal documents including:
- Dust Control Plan
- Infection Control Risk Assessment (ICRA)
- Stormwater Pollution Prevention Plan (SWPPP)
- Fire Safety Plan
- Noise & Vibration Control Plan
"""

import io
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import streamlit as st

# ---------------------------------------------------------------------------
# Page config
# ---------------------------------------------------------------------------
st.set_page_config(
    page_title="Construction Submittal Generator",
    page_icon="🏗️",
    layout="wide",
)

# ---------------------------------------------------------------------------
# Helpers – Word document generation
# ---------------------------------------------------------------------------

def _add_title(doc: Document, title: str):
    p = doc.add_heading(title, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)


def _add_section(doc: Document, heading: str, body: str):
    doc.add_heading(heading, level=2)
    doc.add_paragraph(body)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)


def _add_signature_block(doc: Document):
    doc.add_paragraph("")
    doc.add_paragraph("_" * 40 + "          " + "_" * 20)
    doc.add_paragraph("Authorized Signature                                  Date")
    doc.add_paragraph("")
    doc.add_paragraph("_" * 40 + "          " + "_" * 20)
    doc.add_paragraph("Printed Name                                            Title")


def _doc_to_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Document builders
# ---------------------------------------------------------------------------

def build_dust_control_plan(d: dict) -> bytes:
    doc = Document()
    _add_title(doc, "DUST CONTROL PLAN")

    _add_section(doc, "1. Project Information", "")
    _add_table(doc, ["Field", "Value"], [
        ["Project Name", d["project_name"]],
        ["Project Address", d["project_address"]],
        ["Owner / Developer", d["owner"]],
        ["General Contractor", d["contractor"]],
        ["Project Manager", d["project_manager"]],
        ["Start Date", d["start_date"]],
        ["Estimated Completion", d["end_date"]],
        ["Total Disturbed Area", d["disturbed_area"]],
    ])

    _add_section(doc, "2. Site Description", d["site_description"])

    _add_section(doc, "3. Potential Dust Sources", d["dust_sources"])

    _add_section(doc, "4. Dust Control Measures", "")
    measures = d["control_measures"]
    for m in measures:
        doc.add_paragraph(m, style="List Bullet")

    _add_section(doc, "5. Water Application Schedule", d["water_schedule"])

    _add_section(doc, "6. Monitoring & Inspection", d["monitoring"])

    _add_section(doc, "7. Responsible Personnel", "")
    _add_table(doc, ["Role", "Name", "Phone"], [
        ["Dust Control Supervisor", d["dust_supervisor"], d["dust_supervisor_phone"]],
        ["Site Safety Officer", d["safety_officer"], d["safety_officer_phone"]],
    ])

    _add_section(doc, "8. Emergency Procedures", d["emergency_procedures"])

    _add_section(doc, "9. Regulatory Compliance",
                 f"This plan complies with {d['regulatory_agency']} requirements and all applicable local, "
                 "state, and federal dust and air quality regulations.")

    _add_signature_block(doc)
    return _doc_to_bytes(doc)


def build_infection_control_plan(d: dict) -> bytes:
    doc = Document()
    _add_title(doc, "INFECTION CONTROL RISK ASSESSMENT (ICRA)")

    _add_section(doc, "1. Project Information", "")
    _add_table(doc, ["Field", "Value"], [
        ["Facility Name", d["facility_name"]],
        ["Project Name", d["project_name"]],
        ["Project Location / Floor", d["project_location"]],
        ["General Contractor", d["contractor"]],
        ["Start Date", d["start_date"]],
        ["Estimated Completion", d["end_date"]],
    ])

    _add_section(doc, "2. Construction Activity Type", d["activity_type"])

    _add_section(doc, "3. Risk Group Classification", "")
    _add_table(doc, ["Patient Risk Group", "Areas Affected"], [
        [d["risk_group"], d["areas_affected"]],
    ])

    _add_section(doc, "4. ICRA Class Determination",
                 f"Based on the construction activity type and patient risk group, "
                 f"this project is classified as **ICRA Class {d['icra_class']}**.")

    _add_section(doc, "5. Required Control Measures", "")
    for m in d["control_measures"]:
        doc.add_paragraph(m, style="List Bullet")

    _add_section(doc, "6. Barrier & Containment Description", d["barrier_description"])

    _add_section(doc, "7. HVAC / Ventilation Controls", d["hvac_controls"])

    _add_section(doc, "8. Traffic & Access Routes", d["traffic_routes"])

    _add_section(doc, "9. Monitoring & Inspection Schedule", d["monitoring"])

    _add_section(doc, "10. Responsible Personnel", "")
    _add_table(doc, ["Role", "Name", "Phone"], [
        ["Infection Control Officer", d["ic_officer"], d["ic_officer_phone"]],
        ["Project Manager", d["project_manager"], d["pm_phone"]],
        ["Safety Officer", d["safety_officer"], d["safety_officer_phone"]],
    ])

    _add_section(doc, "11. Approval", "")
    doc.add_paragraph("This ICRA has been reviewed and approved by the Infection Control Department "
                      "and Facilities Management.")
    _add_signature_block(doc)
    return _doc_to_bytes(doc)


def build_swppp(d: dict) -> bytes:
    doc = Document()
    _add_title(doc, "STORMWATER POLLUTION PREVENTION PLAN (SWPPP)")

    _add_section(doc, "1. Project Information", "")
    _add_table(doc, ["Field", "Value"], [
        ["Project Name", d["project_name"]],
        ["Project Address", d["project_address"]],
        ["Owner / Developer", d["owner"]],
        ["General Contractor", d["contractor"]],
        ["NPDES Permit #", d["npdes_permit"]],
        ["Total Site Area", d["total_area"]],
        ["Total Disturbed Area", d["disturbed_area"]],
        ["Start Date", d["start_date"]],
        ["Estimated Completion", d["end_date"]],
    ])

    _add_section(doc, "2. Site Description & Existing Conditions", d["site_description"])

    _add_section(doc, "3. Receiving Waters", d["receiving_waters"])

    _add_section(doc, "4. Potential Pollutant Sources", d["pollutant_sources"])

    _add_section(doc, "5. Erosion & Sediment Controls (BMPs)", "")
    for m in d["erosion_controls"]:
        doc.add_paragraph(m, style="List Bullet")

    _add_section(doc, "6. Good Housekeeping Practices", "")
    for m in d["housekeeping"]:
        doc.add_paragraph(m, style="List Bullet")

    _add_section(doc, "7. Inspection Schedule", d["inspection_schedule"])

    _add_section(doc, "8. Corrective Actions", d["corrective_actions"])

    _add_section(doc, "9. SWPPP Responsible Personnel", "")
    _add_table(doc, ["Role", "Name", "Phone"], [
        ["SWPPP Administrator", d["swppp_admin"], d["swppp_admin_phone"]],
        ["QSP (Qualified SWPPP Practitioner)", d["qsp_name"], d["qsp_phone"]],
    ])

    _add_signature_block(doc)
    return _doc_to_bytes(doc)


def build_fire_safety_plan(d: dict) -> bytes:
    doc = Document()
    _add_title(doc, "FIRE SAFETY PLAN")

    _add_section(doc, "1. Project Information", "")
    _add_table(doc, ["Field", "Value"], [
        ["Project Name", d["project_name"]],
        ["Project Address", d["project_address"]],
        ["General Contractor", d["contractor"]],
        ["Fire Watch Supervisor", d["fire_watch_supervisor"]],
        ["Start Date", d["start_date"]],
        ["Estimated Completion", d["end_date"]],
    ])

    _add_section(doc, "2. Hot Work Activities", d["hot_work"])

    _add_section(doc, "3. Fire Prevention Measures", "")
    for m in d["prevention_measures"]:
        doc.add_paragraph(m, style="List Bullet")

    _add_section(doc, "4. Fire Extinguisher Locations", d["extinguisher_locations"])

    _add_section(doc, "5. Emergency Evacuation Routes", d["evacuation_routes"])

    _add_section(doc, "6. Fire Watch Requirements", d["fire_watch_requirements"])

    _add_section(doc, "7. Emergency Contacts", "")
    _add_table(doc, ["Contact", "Name", "Phone"], [
        ["Fire Department", d["fire_dept"], d["fire_dept_phone"]],
        ["Site Safety Officer", d["safety_officer"], d["safety_officer_phone"]],
        ["Project Manager", d["project_manager"], d["pm_phone"]],
    ])

    _add_signature_block(doc)
    return _doc_to_bytes(doc)


def build_noise_vibration_plan(d: dict) -> bytes:
    doc = Document()
    _add_title(doc, "NOISE & VIBRATION CONTROL PLAN")

    _add_section(doc, "1. Project Information", "")
    _add_table(doc, ["Field", "Value"], [
        ["Project Name", d["project_name"]],
        ["Project Address", d["project_address"]],
        ["General Contractor", d["contractor"]],
        ["Start Date", d["start_date"]],
        ["Estimated Completion", d["end_date"]],
    ])

    _add_section(doc, "2. Sensitive Receptors", d["sensitive_receptors"])

    _add_section(doc, "3. Anticipated Noise Sources", d["noise_sources"])

    _add_section(doc, "4. Permitted Work Hours", d["work_hours"])

    _add_section(doc, "5. Maximum Noise Levels",
                 f"Daytime limit: {d['day_limit']} dBA\nNighttime limit: {d['night_limit']} dBA")

    _add_section(doc, "6. Noise Control Measures", "")
    for m in d["noise_controls"]:
        doc.add_paragraph(m, style="List Bullet")

    _add_section(doc, "7. Vibration Control Measures", "")
    for m in d["vibration_controls"]:
        doc.add_paragraph(m, style="List Bullet")

    _add_section(doc, "8. Monitoring Plan", d["monitoring"])

    _add_section(doc, "9. Community Notification", d["community_notification"])

    _add_section(doc, "10. Complaint Response Procedure", d["complaint_procedure"])

    _add_signature_block(doc)
    return _doc_to_bytes(doc)


# ---------------------------------------------------------------------------
# UI – Sidebar: document type selector
# ---------------------------------------------------------------------------

st.sidebar.title("📋 Document Type")
doc_type = st.sidebar.radio(
    "Select a submittal document to generate:",
    [
        "Dust Control Plan",
        "Infection Control (ICRA)",
        "SWPPP",
        "Fire Safety Plan",
        "Noise & Vibration Control Plan",
    ],
    index=0,
)

st.sidebar.markdown("---")
st.sidebar.info("Fill out the form, then click **Generate Document** at the bottom to download your .docx file.")

# ---------------------------------------------------------------------------
# UI – Main area
# ---------------------------------------------------------------------------

st.title("🏗️ Construction Submittal Document Generator")
st.markdown(f"### {doc_type}")
st.markdown("---")

# ---- Dust Control Plan ----
if doc_type == "Dust Control Plan":
    with st.form("dust_form"):
        st.subheader("Project Information")
        c1, c2 = st.columns(2)
        project_name = c1.text_input("Project Name")
        project_address = c2.text_input("Project Address")
        owner = c1.text_input("Owner / Developer")
        contractor = c2.text_input("General Contractor")
        project_manager = c1.text_input("Project Manager")
        disturbed_area = c2.text_input("Total Disturbed Area (acres / sq ft)")
        start_date = c1.date_input("Start Date", value=datetime.date.today())
        end_date = c2.date_input("Estimated Completion", value=datetime.date.today() + datetime.timedelta(days=180))

        st.subheader("Site & Dust Details")
        site_description = st.text_area("Site Description", placeholder="Describe terrain, soil type, surrounding land use...")
        dust_sources = st.text_area("Potential Dust Sources", placeholder="e.g., excavation, grading, demolition, material stockpiles, haul roads...")

        st.subheader("Control Measures")
        control_options = [
            "Water trucks for haul roads",
            "Soil stabilizers / tackifiers",
            "Covered haul trucks",
            "Wind fencing / barriers",
            "Revegetation of disturbed areas",
            "Gravel at site entry/exit",
            "Street sweeping",
            "Dust suppressant chemicals",
            "Covered stockpiles",
            "Speed limits on unpaved roads",
        ]
        control_measures = st.multiselect("Select applicable measures", control_options, default=control_options[:4])
        custom_measures = st.text_input("Additional measures (comma-separated)", "")
        if custom_measures:
            control_measures += [m.strip() for m in custom_measures.split(",") if m.strip()]

        water_schedule = st.text_area("Water Application Schedule", placeholder="e.g., Every 2 hours on active haul roads; as needed during high-wind events...")

        st.subheader("Monitoring & Compliance")
        monitoring = st.text_area("Monitoring & Inspection Plan", placeholder="Frequency, methods, record-keeping...")
        emergency_procedures = st.text_area("Emergency Procedures", placeholder="Actions during high-wind events, complaints, exceedances...")
        regulatory_agency = st.text_input("Applicable Regulatory Agency", "Local Air Quality Management District")

        st.subheader("Responsible Personnel")
        rc1, rc2 = st.columns(2)
        dust_supervisor = rc1.text_input("Dust Control Supervisor Name")
        dust_supervisor_phone = rc2.text_input("Supervisor Phone")
        safety_officer = rc1.text_input("Site Safety Officer Name")
        safety_officer_phone = rc2.text_input("Safety Officer Phone")

        submitted = st.form_submit_button("🔨 Generate Dust Control Plan", type="primary", use_container_width=True)

    if submitted:
        data = dict(
            project_name=project_name, project_address=project_address, owner=owner,
            contractor=contractor, project_manager=project_manager,
            disturbed_area=disturbed_area,
            start_date=str(start_date), end_date=str(end_date),
            site_description=site_description, dust_sources=dust_sources,
            control_measures=control_measures, water_schedule=water_schedule,
            monitoring=monitoring, emergency_procedures=emergency_procedures,
            regulatory_agency=regulatory_agency,
            dust_supervisor=dust_supervisor, dust_supervisor_phone=dust_supervisor_phone,
            safety_officer=safety_officer, safety_officer_phone=safety_officer_phone,
        )
        doc_bytes = build_dust_control_plan(data)
        st.success("Document generated successfully!")
        st.download_button("⬇️ Download Dust Control Plan (.docx)", doc_bytes,
                           file_name="Dust_Control_Plan.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ---- Infection Control (ICRA) ----
elif doc_type == "Infection Control (ICRA)":
    with st.form("icra_form"):
        st.subheader("Project Information")
        c1, c2 = st.columns(2)
        facility_name = c1.text_input("Healthcare Facility Name")
        project_name = c2.text_input("Project Name")
        project_location = c1.text_input("Project Location / Floor")
        contractor = c2.text_input("General Contractor")
        start_date = c1.date_input("Start Date", value=datetime.date.today())
        end_date = c2.date_input("Estimated Completion", value=datetime.date.today() + datetime.timedelta(days=90))

        st.subheader("Risk Assessment")
        activity_type = st.selectbox("Construction Activity Type", [
            "Type A – Inspection / non-invasive activities",
            "Type B – Small-scale, short-duration work (dust-generating)",
            "Type C – Work generating moderate-high dust, demolition, construction",
            "Type D – Major demolition and construction projects",
        ])
        risk_group = st.selectbox("Patient Risk Group", [
            "Low Risk",
            "Medium Risk",
            "High Risk",
            "Highest Risk (immunocompromised, OR, ICU, etc.)",
        ])
        areas_affected = st.text_input("Patient / Clinical Areas Affected", placeholder="e.g., 3rd floor med-surg, adjacent to ICU...")

        # Auto-determine ICRA class
        activity_idx = ["Type A", "Type B", "Type C", "Type D"].index(activity_type[:6])
        risk_idx = ["Low", "Medium", "High", "Highest"].index(risk_group.split()[0])
        icra_matrix = [
            ["I",  "II",  "II",  "III"],
            ["I",  "II",  "III", "IV"],
            ["I",  "III", "III", "IV"],
            ["II", "III", "IV",  "IV"],
        ]
        icra_class = icra_matrix[activity_idx][risk_idx]
        st.info(f"Auto-determined ICRA Class: **{icra_class}**")

        st.subheader("Control Measures")
        icra_controls = {
            "I": [
                "Minimize dust during work",
                "Replace ceiling tiles immediately",
            ],
            "II": [
                "Wet methods for cutting/sanding",
                "HEPA vacuum before removing barriers",
                "Mist work surfaces to suppress dust",
                "Seal unused doors / grilles",
                "Block and seal return air vents",
                "Wipe surfaces with disinfectant",
            ],
            "III": [
                "Install rigid dust-proof barriers (floor to deck)",
                "Maintain negative air pressure with HEPA filtration",
                "Seal all penetrations in barriers",
                "Anteroom for personnel entry/exit",
                "HEPA vacuum all surfaces before barrier removal",
                "Contain debris in sealed carts for transport",
                "Workers must not travel through patient areas in soiled clothing",
            ],
            "IV": [
                "Install rigid dust-proof barriers (floor to deck) sealed with caulk/tape",
                "Maintain negative air pressure with HEPA filtration (monitor daily)",
                "Anteroom with self-closing doors",
                "HEPA filtered exhaust to outside",
                "All penetrations sealed airtight",
                "Workers shower / change before entering patient areas",
                "Dedicated construction entrance/exit",
                "Daily monitoring of barrier integrity",
                "Air sampling before re-occupancy",
            ],
        }
        default_measures = icra_controls.get(icra_class, [])
        control_measures = st.multiselect("Required control measures", default_measures, default=default_measures)
        extra = st.text_input("Additional measures (comma-separated)", "")
        if extra:
            control_measures += [m.strip() for m in extra.split(",") if m.strip()]

        barrier_description = st.text_area("Barrier & Containment Description",
                                           placeholder="Materials used, dimensions, sealing methods...")
        hvac_controls = st.text_area("HVAC / Ventilation Controls",
                                     placeholder="Negative pressure details, HEPA specs, monitoring method...")
        traffic_routes = st.text_area("Traffic & Access Routes",
                                      placeholder="Worker entry/exit, debris removal paths, patient-free routes...")

        monitoring = st.text_area("Monitoring & Inspection Schedule",
                                  placeholder="Daily barrier checks, weekly air sampling, etc...")

        st.subheader("Responsible Personnel")
        rc1, rc2 = st.columns(2)
        ic_officer = rc1.text_input("Infection Control Officer")
        ic_officer_phone = rc2.text_input("IC Officer Phone")
        project_manager = rc1.text_input("Project Manager")
        pm_phone = rc2.text_input("PM Phone")
        safety_officer = rc1.text_input("Safety Officer")
        safety_officer_phone = rc2.text_input("Safety Officer Phone")

        submitted = st.form_submit_button("🔨 Generate ICRA Document", type="primary", use_container_width=True)

    if submitted:
        data = dict(
            facility_name=facility_name, project_name=project_name,
            project_location=project_location, contractor=contractor,
            start_date=str(start_date), end_date=str(end_date),
            activity_type=activity_type, risk_group=risk_group,
            areas_affected=areas_affected, icra_class=icra_class,
            control_measures=control_measures,
            barrier_description=barrier_description, hvac_controls=hvac_controls,
            traffic_routes=traffic_routes, monitoring=monitoring,
            ic_officer=ic_officer, ic_officer_phone=ic_officer_phone,
            project_manager=project_manager, pm_phone=pm_phone,
            safety_officer=safety_officer, safety_officer_phone=safety_officer_phone,
        )
        doc_bytes = build_infection_control_plan(data)
        st.success("Document generated successfully!")
        st.download_button("⬇️ Download ICRA Document (.docx)", doc_bytes,
                           file_name="ICRA_Plan.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ---- SWPPP ----
elif doc_type == "SWPPP":
    with st.form("swppp_form"):
        st.subheader("Project Information")
        c1, c2 = st.columns(2)
        project_name = c1.text_input("Project Name")
        project_address = c2.text_input("Project Address")
        owner = c1.text_input("Owner / Developer")
        contractor = c2.text_input("General Contractor")
        npdes_permit = c1.text_input("NPDES Permit Number")
        total_area = c2.text_input("Total Site Area")
        disturbed_area = c1.text_input("Total Disturbed Area")
        start_date = c1.date_input("Start Date", value=datetime.date.today())
        end_date = c2.date_input("Estimated Completion", value=datetime.date.today() + datetime.timedelta(days=365))

        st.subheader("Site Conditions")
        site_description = st.text_area("Site Description & Existing Conditions",
                                        placeholder="Topography, soil types, vegetation, drainage patterns...")
        receiving_waters = st.text_input("Receiving Water Bodies",
                                         placeholder="e.g., Oak Creek → San Pedro River")
        pollutant_sources = st.text_area("Potential Pollutant Sources",
                                         placeholder="Sediment, concrete washout, fuel/oil, paint, trash...")

        st.subheader("Best Management Practices (BMPs)")
        erosion_options = [
            "Silt fencing along downslope perimeter",
            "Fiber rolls / wattles",
            "Stabilized construction entrance",
            "Sediment basin / trap",
            "Check dams in drainage swales",
            "Erosion control blankets",
            "Hydroseeding / temporary seeding",
            "Inlet protection (storm drains)",
            "Concrete washout containment area",
            "Earth dikes / berms",
        ]
        erosion_controls = st.multiselect("Erosion & Sediment Controls", erosion_options, default=erosion_options[:5])

        housekeeping_options = [
            "Daily site cleanup of debris",
            "Proper storage of chemicals / materials",
            "Spill prevention & containment kits on-site",
            "Designated fueling area with secondary containment",
            "Covered dumpsters / waste containers",
            "Street sweeping at site exit",
            "Employee training on SWPPP requirements",
        ]
        housekeeping = st.multiselect("Good Housekeeping Practices", housekeeping_options, default=housekeeping_options[:4])

        st.subheader("Inspections & Corrective Actions")
        inspection_schedule = st.text_area("Inspection Schedule",
                                           placeholder="Weekly and within 24 hours of 0.5 inch or greater rainfall...")
        corrective_actions = st.text_area("Corrective Action Procedures",
                                          placeholder="Steps when BMPs fail or discharge observed...")

        st.subheader("Responsible Personnel")
        rc1, rc2 = st.columns(2)
        swppp_admin = rc1.text_input("SWPPP Administrator")
        swppp_admin_phone = rc2.text_input("Admin Phone")
        qsp_name = rc1.text_input("Qualified SWPPP Practitioner (QSP)")
        qsp_phone = rc2.text_input("QSP Phone")

        submitted = st.form_submit_button("🔨 Generate SWPPP", type="primary", use_container_width=True)

    if submitted:
        data = dict(
            project_name=project_name, project_address=project_address,
            owner=owner, contractor=contractor, npdes_permit=npdes_permit,
            total_area=total_area, disturbed_area=disturbed_area,
            start_date=str(start_date), end_date=str(end_date),
            site_description=site_description, receiving_waters=receiving_waters,
            pollutant_sources=pollutant_sources, erosion_controls=erosion_controls,
            housekeeping=housekeeping, inspection_schedule=inspection_schedule,
            corrective_actions=corrective_actions,
            swppp_admin=swppp_admin, swppp_admin_phone=swppp_admin_phone,
            qsp_name=qsp_name, qsp_phone=qsp_phone,
        )
        doc_bytes = build_swppp(data)
        st.success("Document generated successfully!")
        st.download_button("⬇️ Download SWPPP (.docx)", doc_bytes,
                           file_name="SWPPP.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ---- Fire Safety Plan ----
elif doc_type == "Fire Safety Plan":
    with st.form("fire_form"):
        st.subheader("Project Information")
        c1, c2 = st.columns(2)
        project_name = c1.text_input("Project Name")
        project_address = c2.text_input("Project Address")
        contractor = c1.text_input("General Contractor")
        fire_watch_supervisor = c2.text_input("Fire Watch Supervisor")
        start_date = c1.date_input("Start Date", value=datetime.date.today())
        end_date = c2.date_input("Estimated Completion", value=datetime.date.today() + datetime.timedelta(days=180))

        st.subheader("Fire Hazard Assessment")
        hot_work = st.text_area("Hot Work Activities",
                                placeholder="Welding, cutting, brazing, soldering – locations and frequency...")

        prevention_options = [
            "Hot work permit required before any torch/welding work",
            "Fire watch maintained 30 min after hot work completion",
            "Combustibles cleared within 35-foot radius of hot work",
            "Fire extinguishers within 25 feet of hot work",
            "Spark-resistant blankets / shields on combustibles",
            "No hot work near flammable storage areas",
            "Temporary fire alarm / detection in construction areas",
            "Smoking restricted to designated areas only",
            "Flammable material storage in approved cabinets",
            "Daily hot work area inspections",
        ]
        prevention_measures = st.multiselect("Fire Prevention Measures", prevention_options, default=prevention_options[:6])

        extinguisher_locations = st.text_area("Fire Extinguisher Locations",
                                              placeholder="Describe placement: each floor, within 75ft travel distance, near hot work...")
        evacuation_routes = st.text_area("Emergency Evacuation Routes",
                                         placeholder="Primary and secondary routes, assembly point location...")
        fire_watch_requirements = st.text_area("Fire Watch Requirements",
                                               placeholder="When required, duration, trained personnel, communication method...")

        st.subheader("Emergency Contacts")
        rc1, rc2 = st.columns(2)
        fire_dept = rc1.text_input("Fire Department")
        fire_dept_phone = rc2.text_input("Fire Dept Phone", "911")
        safety_officer = rc1.text_input("Site Safety Officer")
        safety_officer_phone = rc2.text_input("Safety Officer Phone")
        project_manager = rc1.text_input("Project Manager")
        pm_phone = rc2.text_input("PM Phone")

        submitted = st.form_submit_button("🔨 Generate Fire Safety Plan", type="primary", use_container_width=True)

    if submitted:
        data = dict(
            project_name=project_name, project_address=project_address,
            contractor=contractor, fire_watch_supervisor=fire_watch_supervisor,
            start_date=str(start_date), end_date=str(end_date),
            hot_work=hot_work, prevention_measures=prevention_measures,
            extinguisher_locations=extinguisher_locations,
            evacuation_routes=evacuation_routes,
            fire_watch_requirements=fire_watch_requirements,
            fire_dept=fire_dept, fire_dept_phone=fire_dept_phone,
            safety_officer=safety_officer, safety_officer_phone=safety_officer_phone,
            project_manager=project_manager, pm_phone=pm_phone,
        )
        doc_bytes = build_fire_safety_plan(data)
        st.success("Document generated successfully!")
        st.download_button("⬇️ Download Fire Safety Plan (.docx)", doc_bytes,
                           file_name="Fire_Safety_Plan.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ---- Noise & Vibration Control Plan ----
elif doc_type == "Noise & Vibration Control Plan":
    with st.form("noise_form"):
        st.subheader("Project Information")
        c1, c2 = st.columns(2)
        project_name = c1.text_input("Project Name")
        project_address = c2.text_input("Project Address")
        contractor = c1.text_input("General Contractor")
        start_date = c1.date_input("Start Date", value=datetime.date.today())
        end_date = c2.date_input("Estimated Completion", value=datetime.date.today() + datetime.timedelta(days=180))

        st.subheader("Noise & Vibration Assessment")
        sensitive_receptors = st.text_area("Sensitive Receptors",
                                           placeholder="Nearby residences, hospitals, schools, churches – with distances...")
        noise_sources = st.text_area("Anticipated Noise Sources",
                                     placeholder="Pile driving, jackhammers, concrete saws, heavy equipment...")
        work_hours = st.text_input("Permitted Construction Hours", "Monday–Friday 7:00 AM – 6:00 PM; Saturday 8:00 AM – 5:00 PM")
        day_limit = st.text_input("Daytime Noise Limit (dBA at property line)", "85")
        night_limit = st.text_input("Nighttime Noise Limit (dBA at property line)", "70")

        st.subheader("Control Measures")
        noise_options = [
            "Equipment fitted with manufacturer-supplied mufflers",
            "Temporary sound barriers / blankets at perimeter",
            "Schedule high-noise activities during least sensitive hours",
            "Use electric equipment where possible (vs. diesel)",
            "Limit simultaneous operation of loud equipment",
            "Maintain equipment to minimize noise from wear",
            "Enclose stationary equipment (generators, compressors)",
            "Pre-notify neighbors before high-impact activities",
        ]
        noise_controls = st.multiselect("Noise Control Measures", noise_options, default=noise_options[:5])

        vibration_options = [
            "Pre-construction survey of adjacent structures",
            "Vibration monitors installed at nearest structures",
            "Use vibratory rollers instead of impact methods where feasible",
            "Limit pile driving to daytime hours only",
            "Use drilled piles instead of driven piles where feasible",
            "Establish vibration thresholds (PPV) for adjacent structures",
            "Cease operations if thresholds exceeded – investigate before resuming",
        ]
        vibration_controls = st.multiselect("Vibration Control Measures", vibration_options, default=vibration_options[:4])

        monitoring = st.text_area("Monitoring Plan",
                                  placeholder="Equipment types, locations, frequency, data logging...")
        community_notification = st.text_area("Community Notification Plan",
                                              placeholder="How and when neighbors are notified of high-noise activities...")
        complaint_procedure = st.text_area("Complaint Response Procedure",
                                           placeholder="Hotline number, response time, investigation steps...")

        submitted = st.form_submit_button("🔨 Generate Noise & Vibration Plan", type="primary", use_container_width=True)

    if submitted:
        data = dict(
            project_name=project_name, project_address=project_address,
            contractor=contractor,
            start_date=str(start_date), end_date=str(end_date),
            sensitive_receptors=sensitive_receptors, noise_sources=noise_sources,
            work_hours=work_hours, day_limit=day_limit, night_limit=night_limit,
            noise_controls=noise_controls, vibration_controls=vibration_controls,
            monitoring=monitoring, community_notification=community_notification,
            complaint_procedure=complaint_procedure,
        )
        doc_bytes = build_noise_vibration_plan(data)
        st.success("Document generated successfully!")
        st.download_button("⬇️ Download Noise & Vibration Plan (.docx)", doc_bytes,
                           file_name="Noise_Vibration_Plan.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ---------------------------------------------------------------------------
st.markdown("---")
st.caption("Construction Submittal Document Generator v1.0 | Documents are generated in .docx format")
